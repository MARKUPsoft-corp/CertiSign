import io  # Importation du module io pour travailler avec les flux de données en mémoire
import os  # Importation du module os pour interagir avec le système de fichiers
import uuid  # Importation du module uuid pour générer des identifiants uniques
import base64  # Importation de base64 pour encoder/décoder en base64
import tempfile  # Importation du module tempfile pour créer des fichiers temporaires
import logging  # Importation de logging pour la journalisation des événements
import datetime  # Importation du module datetime pour la gestion des dates et heures

# Importation des modules de cryptographie nécessaires
from fastapi import FastAPI, File, UploadFile, Form, HTTPException  # FastAPI pour la gestion des requêtes HTTP
from fastapi.responses import FileResponse, JSONResponse  # Pour renvoyer des fichiers et des réponses JSON
from cryptography.hazmat.primitives import hashes, serialization  # Pour les primitives cryptographiques
from cryptography.hazmat.primitives.asymmetric import padding, rsa  # Pour les algorithmes RSA et PSS
from cryptography.hazmat.backends import default_backend  # Pour le backend par défaut de la cryptographie
from cryptography.hazmat.primitives.serialization import pkcs12  # Pour la gestion des fichiers PKCS#12 (PFX)

# Importation des bibliothèques pour la gestion des fichiers PDF et DOCX
import fitz  # PyMuPDF pour les PDF
from docx import Document  # python-docx pour manipuler les fichiers DOCX

# Configuration du logger pour la journalisation des événements
logging.basicConfig(level=logging.INFO)  # Définir le niveau de log à INFO
logger = logging.getLogger(__name__)  # Créer un objet logger pour l'application

app = FastAPI()  # Créer une instance de l'application FastAPI

# Fonction pour charger une clé privée et un certificat à partir d'un fichier PFX (PKCS#12)
def load_private_key_from_pfx(pfx_data: bytes, password: str):
    try:
        # Charger la clé privée, le certificat et les certificats supplémentaires depuis le fichier PFX
        private_key, certificate, additional_certificates = pkcs12.load_key_and_certificates(
            pfx_data,
            password.encode() if password else None,  # Convertir le mot de passe en bytes s'il existe
            backend=default_backend()  # Utiliser le backend par défaut
        )
        
        # Vérifier que la clé privée et le certificat sont bien chargés
        if private_key is None or certificate is None:
            raise ValueError("Clé privée ou certificat introuvable dans le fichier PFX.")
        
        # Vérification de la validité du certificat (date de validité)
        now = datetime.datetime.utcnow()  # Obtenir la date et l'heure actuelles
        if now < certificate.not_valid_before or now > certificate.not_valid_after:
            raise ValueError("Le certificat est expiré ou n'est pas encore valide.")
        
        logger.info("Certificat et clé privée chargés avec succès.")  # Journaliser l'événement
        return private_key, certificate  # Retourner la clé privée et le certificat
    except Exception as e:
        logger.error("Erreur lors du chargement du certificat PFX : %s", e)  # Journaliser l'erreur
        raise HTTPException(status_code=400, detail=f"Erreur lors du chargement du certificat PFX: {str(e)}")  # Lancer une exception HTTP

# Fonction pour signer des données avec la clé privée en utilisant RSA-PSS et SHA256
def sign_bytes(data: bytes, private_key) -> str:
    try:
        # Créer la signature en utilisant RSA-PSS et SHA256
        signature = private_key.sign(
            data,
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),  # MGF1 avec SHA256
                salt_length=padding.PSS.MAX_LENGTH  # Longueur maximale du sel
            ),
            hashes.SHA256()  # Utiliser SHA256 comme fonction de hachage
        )
        return base64.b64encode(signature).decode()  # Retourner la signature encodée en base64
    except Exception as e:
        logger.error("Erreur lors de la signature: %s", e)  # Journaliser l'erreur
        raise HTTPException(status_code=500, detail=f"Erreur lors de la signature: {str(e)}")  # Lancer une exception HTTP

# Fonction pour signer un fichier PDF en ajoutant une annotation de signature
def sign_pdf_file(file_data: bytes, private_key) -> bytes:
    try:
        # Ouvrir le PDF depuis un flux binaire (mémoire)
        doc = fitz.open(stream=file_data, filetype="pdf")
        digital_signature = sign_bytes(file_data, private_key)  # Générer la signature numérique
        annotation_text = f"Document signé numériquement\nSignature: {digital_signature}"  # Texte d'annotation
        
        # Ajouter l'annotation sur la première page du PDF
        first_page = doc[0]  # Première page du document
        first_page.insert_text((50, 50), annotation_text, fontsize=12, color=(1, 0, 0))  # Insérer le texte
        
        # Sauvegarder le PDF signé dans un fichier temporaire
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_file:
            doc.save(tmp_file.name)  # Sauvegarder le document dans un fichier temporaire
            tmp_file.seek(0)  # Revenir au début du fichier
            signed_pdf_bytes = tmp_file.read()  # Lire les données du fichier signé
        doc.close()  # Fermer le document PDF
        return signed_pdf_bytes  # Retourner les données du PDF signé
    except Exception as e:
        logger.error("Erreur lors de la signature du PDF: %s", e)  # Journaliser l'erreur
        raise HTTPException(status_code=500, detail=f"Erreur lors de la signature du PDF: {str(e)}")  # Lancer une exception HTTP

# Fonction pour signer un fichier DOCX en ajoutant des paragraphes de signature
def sign_docx_file(file_data: bytes, private_key) -> bytes:
    try:
        # Charger le document DOCX depuis un flux en mémoire
        document = Document(io.BytesIO(file_data))
        full_text = "\n".join([para.text for para in document.paragraphs])  # Récupérer tout le texte du document
        digital_signature = sign_bytes(full_text.encode(), private_key)  # Générer la signature numérique
        document.add_paragraph("Document signé numériquement")  # Ajouter un paragraphe indiquant que le document est signé
        document.add_paragraph(f"Signature: {digital_signature}")  # Ajouter la signature au document
        
        # Sauvegarder le document signé dans un flux en mémoire
        bio = io.BytesIO()
        document.save(bio)  # Sauvegarder le document dans le flux
        return bio.getvalue()  # Retourner le contenu du document signé
    except Exception as e:
        logger.error("Erreur lors de la signature du DOCX: %s", e)  # Journaliser l'erreur
        raise HTTPException(status_code=500, detail=f"Erreur lors de la signature du DOCX: {str(e)}")  # Lancer une exception HTTP

# Endpoint pour signer un document (TXT, JSON, XML, PDF, DOCX)
@app.post("/sign/")
async def sign_document(
    certificate: UploadFile = File(...),  # Fichier PFX contenant le certificat et la clé privée
    password: str = Form(...),  # Mot de passe pour déverrouiller le fichier PFX
    file: UploadFile = File(...)  # Fichier à signer (peut être un PDF, DOCX, etc.)
):
    try:
        file_data = await file.read()  # Lire les données du fichier à signer
        pfx_data = await certificate.read()  # Lire les données du fichier PFX
        private_key, certificate_obj = load_private_key_from_pfx(pfx_data, password)  # Charger la clé privée et le certificat
        
        file_extension = file.filename.split(".")[-1].lower()  # Récupérer l'extension du fichier

        # Traitement pour les fichiers texte
        if file_extension in ["txt", "json", "xml"]:
            digital_signature = sign_bytes(file_data, private_key)  # Signer les données
            return {"filename": file.filename, "signature": digital_signature}  # Retourner la signature en réponse

        # Traitement pour les fichiers PDF
        elif file_extension == "pdf":
            signed_pdf_bytes = sign_pdf_file(file_data, private_key)  # Signer le fichier PDF
            tmp_path = f"signed_{uuid.uuid4().hex}.pdf"  # Générer un nom de fichier temporaire
            with open(tmp_path, "wb") as f_out:
                f_out.write(signed_pdf_bytes)  # Sauvegarder le PDF signé
            response = FileResponse(tmp_path, media_type="application/pdf", filename=os.path.basename(tmp_path))  # Retourner le fichier signé
            return response

        # Traitement pour les fichiers DOCX
        elif file_extension == "docx":
            signed_docx_bytes = sign_docx_file(file_data, private_key)  # Signer le fichier DOCX
            tmp_path = f"signed_{uuid.uuid4().hex}.docx"  # Générer un nom de fichier temporaire
            with open(tmp_path, "wb") as f_out:
                f_out.write(signed_docx_bytes)  # Sauvegarder le fichier DOCX signé
            response = FileResponse(tmp_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=os.path.basename(tmp_path))  # Retourner le fichier signé
            return response

        else:
            logger.warning("Format non supporté: %s", file_extension)  # Journaliser l'avertissement pour un format non supporté
            raise HTTPException(status_code=400, detail="Format non supporté. Utilisez TXT, JSON, XML, PDF ou DOCX.")  # Lancer une exception HTTP

    except Exception as e:
        logger.error("Erreur dans l'endpoint /sign/ : %s", e)  # Journaliser l'erreur
        raise  # Relancer l'exception

# Endpoint pour vérifier la signature d'un document
@app.post("/verify/")
async def verify_document(
    certificate: UploadFile = File(...),  # Fichier PFX contenant le certificat
    password: str = Form(...),  # Mot de passe pour déverrouiller le fichier PFX
    file: UploadFile = File(...),  # Fichier à vérifier
    signature: str = Form(...)  # Signature à vérifier
):
    try:
        file_data = await file.read()  # Lire les données du fichier
        pfx_data = await certificate.read()  # Lire les données du fichier PFX
        private_key, certificate_obj = load_private_key_from_pfx(pfx_data, password)  # Charger la clé privée et le certificat
        public_key = certificate_obj.public_key()  # Obtenir la clé publique du certificat
        
        file_extension = file.filename.split(".")[-1].lower()  # Récupérer l'extension du fichier

        # Vérification de la signature pour les formats texte
        if file_extension in ["txt", "json", "xml"]:
            try:
                # Vérifier la signature avec la clé publique
                public_key.verify(
                    base64.b64decode(signature),  # Décoder la signature
                    file_data,  # Données du fichier
                    padding.PSS(  # Utiliser le padding PSS
                        mgf=padding.MGF1(hashes.SHA256()),  # MGF1 avec SHA256
                        salt_length=padding.PSS.MAX_LENGTH  # Longueur du sel
                    ),
                    hashes.SHA256()  # Fonction de hachage SHA256
                )
                return JSONResponse(content={"filename": file.filename, "verified": True})  # Retourner un résultat positif
            except Exception as e:
                logger.info("Vérification échouée: %s", e)  # Journaliser l'échec
                return JSONResponse(content={"filename": file.filename, "verified": False})  # Retourner un résultat négatif
        
        else:
            raise HTTPException(status_code=400, detail="Format non supporté. Utilisez TXT, JSON ou XML.")  # Lancer une exception pour format non supporté

    except Exception as e:
        logger.error("Erreur dans l'endpoint /verify/ : %s", e)  # Journaliser l'erreur
        raise HTTPException(status_code=500, detail=f"Erreur lors de la vérification: {str(e)}")  # Lancer une exception HTTP
